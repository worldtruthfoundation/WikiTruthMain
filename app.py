import os
import logging
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file
from flask_session import Session
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import DeclarativeBase
from werkzeug.middleware.proxy_fix import ProxyFix
from datetime import datetime
import json
import threading
import io
from docx import Document
from docx.shared import Inches

from wikipedia_api import WikipediaAPI
from openai_service import OpenAIService


# Set up logging
logging.basicConfig(level=logging.DEBUG)

class Base(DeclarativeBase):
    pass

db = SQLAlchemy(model_class=Base)



# Create the app
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "dev-secret-key-change-in-production")
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

app.config['SESSION_TYPE'] = 'filesystem'  # хранить сессии на диске (можно также 'redis', если захочешь)
Session(app)


# Configure the database
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL", "sqlite:///wikitruth.db")
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}

# Initialize extensions
db.init_app(app)

# Initialize services
wiki_api = WikipediaAPI()
openai_service = OpenAIService()

with app.app_context():
    import models
    db.create_all()

@app.route('/')
def index():
    """Homepage with search functionality"""
    languages = wiki_api.get_supported_languages()
    return render_template('index.html', languages=languages)

@app.route('/api/search')
def api_search():
    """API endpoint for article search suggestions"""
    query = request.args.get('q', '').strip()
    language = request.args.get('lang', 'en')
    
    if len(query) < 2:
        return jsonify([])
    
    try:
        suggestions = wiki_api.search_articles(query, language, limit=10)
        return jsonify(suggestions)
    except Exception as e:
        logging.error(f"Search error: {e}")
        return jsonify([])

@app.route('/article/<language>/<path:title>')
def article_selection(language, title):
    """Language selection page for a specific article"""
    try:
        # Decode URL-encoded title
        from urllib.parse import unquote
        decoded_title = unquote(title)
        
        # Get article info
        article_info = wiki_api.get_article_info(decoded_title, language)
        if not article_info:
            flash('Article not found.', 'error')
            return redirect(url_for('index'))
        
        # Get available language versions
        language_versions = wiki_api.get_language_links(decoded_title, language)
        
        # Get supported output languages
        output_languages = wiki_api.get_supported_languages()
        
        return render_template(
            'language_selection.html',
            article_info=article_info,
            language_versions=language_versions,
            output_languages=output_languages,
            selected_language=language
        )
    
    except Exception as e:
        logging.error(f"Error in article_selection: {e}")
        flash('Error loading article information.', 'error')
        return redirect(url_for('index'))
@app.route('/compare', methods=['POST'])
def compare_articles():
    try:
        # Очистить предыдущий результат сравнения (ТОЛЬКО ЗДЕСЬ)
        session.pop('comparison_result', None)

        # Получить данные формы
        languages = request.form.getlist('languages')
        output_language = request.form.get('output_language', 'en')
        mode = request.form.get('mode', 'normal')

        if len(languages) < 2:
            flash('Please select at least 2 languages for comparison.', 'error')
            return redirect(request.referrer or url_for('index'))
        if len(languages) > 5:
            flash('Please select no more than 5 languages for comparison.', 'error')
            return redirect(request.referrer or url_for('index'))

        referer = request.referrer or ''
        base_language, article_title = None, None
        if '/article/' in referer:
            from urllib.parse import unquote
            parts = referer.split('/article/')[1].split('/')
            base_language = parts[0]
            article_title = unquote('/'.join(parts[1:]))
        else:
            flash('Invalid article reference.', 'error')
            return redirect(url_for('index'))

        session['comparison_params'] = {
            'languages': languages,
            'output_language': output_language,
            'mode': mode,
            'base_language': base_language,
            'article_title': article_title,
            'timestamp': datetime.now().isoformat()
        }
        session.modified = True

        return render_template('loading.html')
    except Exception as e:
        logging.error(f"Error starting comparison: {e}")
        flash('Error starting comparison process.', 'error')
        return redirect(url_for('index'))


@app.route('/perform_comparison')
def perform_comparison():
    """
    AJAX endpoint to perform the actual comparison.
    """
    try:
        params = session.get('comparison_params')
        if not params:
            return jsonify({'error': 'No comparison parameters found'}), 400

        language_versions = wiki_api.get_language_links(params['article_title'], params['base_language'])
        language_title_map = {ver['lang']: ver['title'] for ver in language_versions}
        article_requests = [
            {'language': lang, 'title': language_title_map.get(lang, params['article_title'])}
            for lang in params['languages']
        ]
        articles = wiki_api.fetch_articles_parallel(article_requests)

        if len(articles) < 2:
            return jsonify({'error': 'Could not fetch enough articles for comparison'}), 400

        comparison_result = openai_service.compare_articles(
            articles, params['output_language'], params['mode']
        )
        session['comparison_result'] = {'result': comparison_result, 'params': params}
        session.modified = True

        return jsonify({'success': True, 'redirect': url_for('comparison_result')})

    except Exception as e:
        logging.error(f"Error performing comparison: {e}")
        session.modified = True
        return jsonify({'error': f'Comparison failed: {str(e)}'}), 500


@app.route('/result')
def comparison_result():
    """Display comparison results"""
    result_data = session.get('comparison_result')
    if not result_data:
        flash('No comparison result found.', 'error')
        return redirect(url_for('index'))
    
    return render_template(
        'comparison.html',
        result=result_data['result'],
        params=result_data['params']
    )

@app.route('/export/docx')
def export_docx():
    """Export comparison result as Word document"""
    try:
        result_data = session.get('comparison_result')
        if not result_data:
            flash('No comparison result found.', 'error')
            return redirect(url_for('index'))
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Wikipedia Article Comparison', 0)
        
        # Add metadata
        doc.add_heading('Comparison Details', level=1)
        details = doc.add_paragraph()
        details.add_run('Article: ').bold = True
        details.add_run(result_data['params']['article_title'])
        details.add_run('\nLanguages: ').bold = True
        details.add_run(', '.join(result_data['params']['languages']))
        details.add_run('\nOutput Language: ').bold = True
        details.add_run(result_data['params']['output_language'])
        details.add_run('\nMode: ').bold = True
        details.add_run(result_data['params']['mode'].title())
        details.add_run('\nGenerated: ').bold = True
        details.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        
        # Add comparison result
        doc.add_heading('Comparison Result', level=1)
        doc.add_paragraph(result_data['result'])
        
        # Add disclaimer
        doc.add_heading('Disclaimer', level=1)
        disclaimer = doc.add_paragraph()
        disclaimer.add_run('Important: ').bold = True
        disclaimer.add_run(
            'This comparison was generated by artificial intelligence (GPT-4) and should be used for '
            'educational and informational purposes only. While we strive for accuracy, AI-generated '
            'content may contain errors or biases. Please verify important information through multiple sources.'
        )
        
        # Save to memory
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        filename = f"wikipedia_comparison_{result_data['params']['article_title'].replace(' ', '_')}.docx"
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        logging.error(f"Error exporting document: {e}")
        flash('Error generating document export.', 'error')
        return redirect(url_for('comparison_result'))

@app.errorhandler(404)
def not_found_error(error):
    return render_template('base.html'), 404

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('base.html'), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
