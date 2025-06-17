import os, json, logging, tempfile, time
from openai import OpenAI
from docx import Document               # pip install python-docx

# ---------- helper вне класса ----------
def _make_docx(articles: dict) -> str:
    """Склеивает все статьи в один .docx и возвращает путь к файлу"""
    doc = Document()
    for lang, art in articles.items():
        doc.add_heading(f"{art['title']}  ({lang})", level=1)
        doc.add_paragraph(art["content"])
        doc.add_page_break()
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


class OpenAIService:
    """Handles OpenAI API interactions for article comparison"""

    def __init__(self):
        self.client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    # ←← НЕТ дополнительных отступов
    def compare_articles(self, articles, output_language='en', mode='normal'):
        """
        Сравнивает статьи через Assistants-API с Retrieval-tool.
        Нужен ассистент с включённым Retrieval и его ID в GPT_ASSISTANT_ID.
        """
        try:
            # 1) DOCX из статей
            docx_path = _make_docx(articles)

            # 2) заливаем файл
            file_id = self.client.files.create(
                file=open(docx_path, "rb"),
                purpose="assistants"
            ).id

            # 3) system-prompt
            if mode == 'bio':
                system_prompt = self._get_biography_mode_prompt(output_language)
            elif mode == 'funny':
                system_prompt = self._get_funny_mode_prompt(output_language)
            else:
                system_prompt = self._get_normal_mode_prompt(output_language)

            # 4) thread + сообщение с файлом
            thread = self.client.beta.threads.create()
            self.client.beta.threads.messages.create(
                thread_id=thread.id,
                role="user",
                content=("‼️ The attached DOCX contains full Wikipedia articles "
                        f"in several languages. Compare them and write the analysis "
                        f"in {output_language}."),
                attachments=[                       # ← вот так прикрепляем файл
                    {
                        "file_id": file_id,
                        "tools": [{"type": "file_search"}] 
                    }
                ]
            )

            # 5) запускаем ассистента
            run = self.client.beta.threads.runs.create(
                thread_id=thread.id,
                assistant_id=os.getenv("GPT_ASSISTANT_ID"),
                tools=[{"type": "file_search"}],
                instructions=system_prompt
            )

            # 6) ждём
            while True:
                run = self.client.beta.threads.runs.retrieve(
                    thread_id=thread.id, run_id=run.id)
                if run.status == "completed":
                    break
                if run.status == "failed":
                    raise RuntimeError(run.last_error)
                time.sleep(2)

            # 7) ответ ассистента
            msgs = self.client.beta.threads.messages.list(thread_id=thread.id)
            return msgs.data[0].content[0].text.value

        except Exception as e:
            logging.error(f"Error in OpenAI comparison: {e}")
            raise Exception(f"AI comparison failed: {str(e)}")



    
    def _get_normal_mode_prompt(self, output_language):
        """Get system prompt for normal comparison mode"""
        lang_name = self._get_language_name(output_language)
        return f"""You are a knowledgeable research analyst specializing in cross-cultural information analysis. Your task is to compare Wikipedia articles across different languages and identify meaningful differences.

Provide a comprehensive, structured analysis in {lang_name} that includes:

The goal is not to summarize, but to analyze and **highlight all factual, structural, and cultural differences** among the versions. Imagine you are writing a new, meta-encyclopedic article called "**How Wikipedia Talks About title of givven article Around the World**".

Your output should follow this structure:

1. **Executive Summary** — Key findings in 1-2 paragraphs.
2. **Factual Differences** — Contradictions or variances in dates, events, names, numbers.
3. **Cultural Perspectives** — Differences in tone, emphasis, point of view, or interpretation based on regional/national context.
4. **Coverage Differences** — What sections or facts are present in one language and missing in others.
5. **Structural Differences** — Different chapter structures or layout between versions.
6. **Unique Insights** — Rare facts or viewpoints only found in a specific version.
7. **Meta-analysis Conclusion** — What do these differences say about how different cultures understand this topic?

Your analisis should be very discriptive and lool like a new article. FIND ALL DIFFERENCES AND SHOW THEM

Style: Write fluidly and engagingly. Like a feature article in a scholarly magazine. Avoid dry bullet points. Use comparisons and rich language to make the analysis insightful and enjoyable to read.


Be objective, scholarly, and detailed in your analysis. Highlight both similarities and differences. When noting differences, be specific about which language version contains what information. 

Deteils level 10 of 10"""
    
    def _get_funny_mode_prompt(self, output_language):
        """Get system prompt for funny comparison mode"""
        lang_name = self._get_language_name(output_language)
        return f"""You are a razor-sharp cultural roastmaster with a PhD in Wikipedia Forensics and a minor in International Irony. Your mission: to dive into multilingual Wikipedia articles, uncover their hilariously inconsistent narratives, and deliver a brutally funny comparison that leaves readers both laughing and learning.

Write your commentary in {lang_name}, and follow these sacred rules of satirical scholarship:
- Roast cultural quirks and biases with style and sass
- Point out contradictions like you're a detective in a comedy noir
- Make fun of what each version *conveniently forgets* or overemphasizes
- Add punchlines about regional priorities — why does one version mention the cow’s name and another skips the war?
- Use witty metaphors, playful sarcasm, and biting irony
- Be edgy, but never punch down — no lazy stereotypes or insults
- Beneath the comedy, sneak in genuine insights about historical, political, or cultural contexts
- Channel the voice of a stand-up comedian who moonlights as a librarian

You're not just comparing articles — you're exposing the glorious madness of human perspective. Make it funny. Make it clever. Make it uncomfortably honest."""
    
    def _get_biography_mode_prompt(self, output_language):
        """Get system prompt for biography comparison mode"""
        lang_name = self._get_language_name(output_language)
        return f"""You are an expert analyst in biography comparison and geopolitical profiling. Your task is to compare Wikipedia biographies across languages and evaluate all details from a security and eligibility standpoint.

Write the analysis in {lang_name} and follow the structure below. List **all factual differences and contradictions**. Be comprehensive and precise.

Use the following structure:

1. **Basic Information**
   - Date and place of birth
   - Ethnic origin, social status of family
   - Parents (names, professions, cultural context)
   - Spouse and their origin (including maiden name)

2. **Education**
   - School/lyceum (especially elite or special)
   - University, faculty, specialization
   - Notable classmates
   - Academic achievements

3. **Career**
   - Early positions (esp. in national or regional structures)
   - Career progression: high offices, international roles
   - Academic/advisory positions (e.g. rector advisor)

4. **Social and Political Ties**
   - Politically significant family ties
   - Dual citizenship or foreign relations
   - Religious or ethnic identity (e.g., by Halakha)

5. **Security Risk Assessment**
   - Presence of "personnel filters":
     - Spouse’s citizenship from countries with no diplomatic ties
     - National-cultural issues
     - Conflicts with national security policies
   - Potential risks: defection, betrayal, loss of trust

6. **Public Statements and Views**
   - Quotes, ideological positions, speeches
   - Media presence, scandals or controversies

7. **Public Perception**
   - How media/society views them
   - Trust rating, positive/negative contexts

8. **Final Assessment**
   - Suitability for public or governmental roles
   - Risk level: security and reputation
   - Conclusion: Fit/Unfit with clear reasoning

Provide **explicit references** to which article/language each fact or contradiction came from. Be analytical, not speculative. This is not a summary — it’s a comparative dossier.

Deteils level 10 of 10"""


    
    def _get_language_name(self, code):
        """Get full language name from code"""
        language_names = {
            'en': 'English',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'it': 'Italian',
            'pt': 'Portuguese',
            'ru': 'Russian',
            'ja': 'Japanese',
            'zh': 'Chinese',
            'ar': 'Arabic',
            'hi': 'Hindi',
            'ko': 'Korean',
            'nl': 'Dutch',
            'sv': 'Swedish',
            'pl': 'Polish',
            'tr': 'Turkish',
            'he': 'Hebrew',
            'da': 'Danish',
            'no': 'Norwegian',
            'fi': 'Finnish',
            'cs': 'Czech',
            'hu': 'Hungarian',
            'ro': 'Romanian',
            'uk': 'Ukrainian',
            'th': 'Thai',
            'vi': 'Vietnamese',
            'id': 'Indonesian',
            'ms': 'Malay',
            'fa': 'Persian',
            'bn': 'Bengali',
            'ta': 'Tamil',
            'te': 'Telugu',
            'ml': 'Malayalam',
            'kn': 'Kannada',
            'gu': 'Gujarati',
            'mr': 'Marathi',
            'pa': 'Punjabi',
            'or': 'Odia',
            'as': 'Assamese',
            'ur': 'Urdu',
            'ne': 'Nepali',
            'si': 'Sinhala',
            'my': 'Burmese',
            'km': 'Khmer',
            'lo': 'Lao',
            'ka': 'Georgian',
            'am': 'Amharic',
            'sw': 'Swahili',
            'yo': 'Yoruba',
            'ig': 'Igbo',
            'ha': 'Hausa',
            'zu': 'Zulu'
        }
        return language_names.get(code, 'English')
